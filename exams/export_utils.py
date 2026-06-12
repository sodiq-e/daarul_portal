"""
Exam Export Utilities

Handles PDF and DOCX export for exam papers with formatting preservation.
"""

from io import BytesIO
from urllib.parse import urlparse
from urllib.request import urlopen
from urllib.error import URLError, HTTPError
from django.conf import settings
from django.http import FileResponse
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
from html.parser import HTMLParser


class HTMLToDocXConverter(HTMLParser):
    """Convert HTML content to python-docx compatible format"""
    
    def __init__(self):
        super().__init__()
        self.elements = []
        self.current_style = {}
        self.in_list = False
        self.list_type = None
        self.list_depth = 0
        
    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        
        if tag == 'p':
            self.elements.append(('paragraph', {'content': [], 'align': attrs_dict.get('align', 'left')}))
        elif tag == 'br':
            self.elements.append(('break', {}))
        elif tag in ['b', 'strong']:
            self.current_style['bold'] = True
        elif tag in ['i', 'em']:
            self.current_style['italic'] = True
        elif tag == 'u':
            self.current_style['underline'] = True
        elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag[1])
            self.elements.append(('heading', {'level': level, 'content': []}))
        elif tag == 'ul':
            self.list_type = 'bullet'
            self.in_list = True
            self.list_depth += 1
        elif tag == 'ol':
            self.list_type = 'number'
            self.in_list = True
            self.list_depth += 1
        elif tag == 'li':
            if self.in_list:
                self.elements.append(('list_item', {'content': [], 'style': self.list_type, 'depth': self.list_depth}))
        elif tag == 'table':
            self.elements.append(('table', {'rows': []}))
        elif tag == 'tr':
            if self.elements and self.elements[-1][0] == 'table':
                self.elements[-1][1]['rows'].append({'cells': []})
        elif tag in ['td', 'th']:
            if self.elements and self.elements[-1][0] == 'table' and self.elements[-1][1]['rows']:
                self.elements[-1][1]['rows'][-1]['cells'].append({'content': [], 'is_header': tag == 'th'})
        elif tag == 'img':
            src = attrs_dict.get('src', '')
            self.elements.append(('image', {'src': src}))
        elif tag == 'sub':
            self.current_style['subscript'] = True
        elif tag == 'sup':
            self.current_style['superscript'] = True
            
    def handle_endtag(self, tag):
        if tag in ['ul', 'ol']:
            self.list_depth = max(0, self.list_depth - 1)
            if self.list_depth == 0:
                self.in_list = False
        elif tag in ['b', 'strong']:
            self.current_style['bold'] = False
        elif tag in ['i', 'em']:
            self.current_style['italic'] = False
        elif tag == 'u':
            self.current_style['underline'] = False
        elif tag == 'sub':
            self.current_style['subscript'] = False
        elif tag == 'sup':
            self.current_style['superscript'] = False
            
    def handle_data(self, data):
        if data.strip():
            self.elements.append(('text', {'content': data, 'style': self.current_style.copy()}))


def strip_html_tags(html_content):
    """Remove HTML tags from content, preserving text"""
    if not html_content:
        return ''
    
    # Remove script and style elements
    clean = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
    clean = re.sub(r'<style[^>]*>.*?</style>', '', clean, flags=re.DOTALL)
    
    # Remove HTML tags
    clean = re.sub(r'<[^>]+>', '', clean)
    
    # Decode HTML entities and clean up whitespace
    clean = clean.replace('&nbsp;', ' ')
    clean = clean.replace('&lt;', '<')
    clean = clean.replace('&gt;', '>')
    clean = clean.replace('&amp;', '&')
    clean = ' '.join(clean.split())
    
    return clean


def set_cell_border(cell, **kwargs):
    """Set cell borders for tables"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '12')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    
    tcPr.append(tcBorders)


def fetch_image_bytes(src):
    if not src:
        return None
    src = src.strip()
    if src.startswith('//'):
        src = 'https:' + src
    parsed = urlparse(src)
    try:
        if parsed.scheme in ('http', 'https'):
            return urlopen(src, timeout=10).read()
        if src.startswith('/'):
            local_path = os.path.join(settings.BASE_DIR, src.lstrip('/'))
            if os.path.exists(local_path):
                with open(local_path, 'rb') as f:
                    return f.read()
    except (URLError, HTTPError, OSError):
        return None
    return None


def add_html_to_doc(doc, html, style_name=None, left_indent=0):
    if not html:
        return

    # Simple table support
    table_match = re.search(r'<table.*?>(.*?)</table>', html, flags=re.S)
    if table_match:
        table_html = table_match.group(1)
        rows = re.findall(r'<tr.*?>(.*?)</tr>', table_html, flags=re.S)
        if rows:
            first_row_cells = re.findall(r'<t[dh].*?>(.*?)</t[dh]>', rows[0], flags=re.S)
            table = doc.add_table(rows=len(rows), cols=len(first_row_cells))
            table.style = 'Table Grid'
            for row_idx, row_html in enumerate(rows):
                cells = re.findall(r'<t[dh].*?>(.*?)</t[dh]>', row_html, flags=re.S)
                for col_idx, cell_html in enumerate(cells):
                    text = strip_html_tags(cell_html)
                    cell = table.cell(row_idx, col_idx)
                    cell.text = text
                    if row_idx == 0:
                        set_cell_border(cell, top=True, left=True, bottom=True, right=True)
            return

    parts = re.split(r'(<img[^>]+>)', html, flags=re.I)
    for part in parts:
        if not part:
            continue
        if part.lower().startswith('<img'):
            src_match = re.search(r'src=["\']([^"\']+)["\']', part)
            if src_match:
                image_bytes = fetch_image_bytes(src_match.group(1))
                if image_bytes:
                    image_stream = BytesIO(image_bytes)
                    try:
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(image_stream, width=Inches(4))
                    except Exception:
                        doc.add_paragraph('[Image could not be loaded]')
            continue
        text = strip_html_tags(part)
        if not text:
            continue
        paragraph = doc.add_paragraph(text)
        if style_name:
            paragraph.style = style_name
        if left_indent:
            paragraph.paragraph_format.left_indent = Inches(left_indent)


def export_exam_to_docx(exam_paper, include_answers=False, include_marks=True):
    """
    Export exam paper to DOCX format
    
    Args:
        exam_paper: ExamPaper instance
        include_answers: Include teacher guide and answers
        include_marks: Include mark allocations
        
    Returns:
        BytesIO object with DOCX content
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{exam_paper.subject.name}")
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    # Add exam info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = f"{exam_paper.school_class} | {exam_paper.term} | {exam_paper.academic_session}"
    info_run = info.add_run(info_text)
    info_run.font.size = Pt(10)
    
    # Add duration and marks
    duration_marks = doc.add_paragraph()
    duration_marks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dm_text = f"Duration: {exam_paper.duration or 'N/A'} | Total Marks: {exam_paper.total_marks}"
    dm_run = duration_marks.add_run(dm_text)
    dm_run.font.size = Pt(10)
    
    # Add instructions
    if exam_paper.instructions:
        doc.add_heading('Instructions', level=2)
        add_html_to_doc(doc, exam_paper.instructions)
    
    doc.add_paragraph()  # Spacing
    
    # Add sections and questions
    for section_idx, section in enumerate(exam_paper.sections.all(), 1):
        # Section heading
        section_heading = doc.add_heading(f"{section.get_section_type_display()}", level=2)
        
        # Section instruction
        if section.instruction:
            add_html_to_doc(doc, section.instruction, style_name='List Bullet')
        
        if include_marks and section.marks_allocation:
            doc.add_paragraph(f"Total Marks for this Section: {section.marks_allocation}")
        
        doc.add_paragraph()  # Spacing
        
        # Questions
        for q_idx, question in enumerate(section.questions.all(), 1):
            q_para = doc.add_paragraph(style='List Number')
            q_num_run = q_para.add_run(f"{q_idx}. ")
            q_num_run.bold = True
            
            if include_marks and question.marks:
                marks_para = doc.add_paragraph(f"[{question.marks} marks]")
                marks_para.paragraph_format.left_indent = Inches(0.5)
            
            # Question content with formatting and images
            if question.question_text:
                add_html_to_doc(doc, question.question_text, left_indent=0)
            
            if include_answers and question.teacher_guide:
                doc.add_paragraph("Teacher Guide:", style='Heading 3')
                add_html_to_doc(doc, question.teacher_guide, left_indent=0.5)
            
            if question.resource_notes:
                doc.add_paragraph("Supplementary Resources:", style='Heading 4')
                add_html_to_doc(doc, question.resource_notes, left_indent=0.5)
            
            if question.options.exists():
                for option in question.options.all():
                    option_para = doc.add_paragraph(style='List Bullet 2')
                    option_text = strip_html_tags(option.option_text)
                    option_para.add_run(f"{option.option_label}. {option_text}")
            
            doc.add_paragraph()  # Spacing between questions
        
        doc.add_page_break()
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output


def export_exam_to_pdf_html(exam_paper, include_answers=False, include_marks=True):
    """
    Generate HTML suitable for PDF conversion
    
    Args:
        exam_paper: ExamPaper instance
        include_answers: Include teacher guide and answers
        include_marks: Include mark allocations
        
    Returns:
        HTML string
    """
    
    sections_data = []
    
    for section in exam_paper.sections.all():
        questions_data = []
        
        for q_idx, question in enumerate(section.questions.all(), 1):
            question_data = {
                'number': q_idx,
                'text': question.question_text,
                'marks': question.marks if include_marks else None,
                'options': [],
                'teacher_guide': question.teacher_guide if include_answers else None,
            }
            
            for option in question.options.all():
                question_data['options'].append({
                    'label': option.option_label,
                    'text': option.option_text,
                })
            
            questions_data.append(question_data)
        
        section_data = {
            'type': section.get_section_type_display(),
            'title': section.title,
            'instruction': section.instruction,
            'marks_allocation': section.marks_allocation if include_marks else None,
            'questions': questions_data,
        }
        
        sections_data.append(section_data)
    
    context = {
        'exam': exam_paper,
        'sections': sections_data,
        'include_marks': include_marks,
        'include_answers': include_answers,
    }
    
    # You'll need to create an exam_export.html template
    # For now, returning structured data for template rendering
    return context


# Template for PDF export (to be placed in templates/exams/exam_export.html)
EXAM_EXPORT_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {
            font-family: Arial, sans-serif;
            color: #000;
            margin: 2cm;
            background: white;
        }
        .header {
            text-align: center;
            margin-bottom: 2cm;
            border-bottom: 2px solid #000;
            padding-bottom: 1cm;
        }
        .title {
            font-size: 20px;
            font-weight: bold;
            margin-bottom: 10px;
        }
        .exam-info {
            font-size: 12px;
            margin-bottom: 5px;
        }
        .instructions {
            background: #f9f9f9;
            padding: 1cm;
            margin: 1cm 0;
            border-left: 4px solid #ccc;
        }
        .section {
            page-break-inside: avoid;
            margin: 1cm 0;
        }
        .section-title {
            font-size: 16px;
            font-weight: bold;
            margin-bottom: 0.5cm;
            text-decoration: underline;
        }
        .question {
            margin: 0.5cm 0;
            page-break-inside: avoid;
        }
        .question-number {
            font-weight: bold;
            display: inline-block;
            margin-right: 5px;
        }
        .marks {
            float: right;
            font-weight: bold;
            color: #666;
        }
        .option {
            margin-left: 2cm;
            margin-bottom: 3px;
        }
        .teacher-guide {
            background: #f0f0f0;
            padding: 0.5cm;
            margin: 0.5cm 0 0.5cm 2cm;
            border-left: 3px solid #999;
            font-size: 11px;
        }
        .teacher-guide-label {
            font-weight: bold;
            font-size: 10px;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 0.5cm 0;
        }
        th, td {
            border: 1px solid #000;
            padding: 5px;
            text-align: left;
        }
        th {
            background: #f0f0f0;
            font-weight: bold;
        }
    </style>
</head>
<body>
    <div class="header">
        <div class="title">{{ exam.subject.name }}</div>
        <div class="exam-info">{{ exam.school_class }} | {{ exam.term }} | {{ exam.academic_session }}</div>
        <div class="exam-info">Duration: {{ exam.duration|default:"N/A" }} | Total Marks: {{ exam.total_marks }}</div>
    </div>
    
    {% if exam.instructions %}
    <div class="instructions">
        <strong>Instructions:</strong>
        {{ exam.instructions|safe }}
    </div>
    {% endif %}
    
    {% for section in sections %}
    <div class="section">
        <div class="section-title">{{ section.type }}</div>
        {% if section.instruction %}
        <p>{{ section.instruction|safe }}</p>
        {% endif %}
        {% if section.marks_allocation %}
        <p><strong>Total Marks: {{ section.marks_allocation }}</strong></p>
        {% endif %}
        
        {% for question in section.questions %}
        <div class="question">
            <div>
                <span class="question-number">{{ question.number }}.</span>
                {{ question.text|safe }}
                {% if question.marks %}
                <span class="marks">[{{ question.marks }} marks]</span>
                {% endif %}
            </div>
            
            {% if question.options %}
            {% for option in question.options %}
            <div class="option">
                <strong>{{ option.label }}.</strong> {{ option.text|safe }}
            </div>
            {% endfor %}
            {% endif %}
            
            {% if question.teacher_guide %}
            <div class="teacher-guide">
                <div class="teacher-guide-label">Teacher Guide:</div>
                {{ question.teacher_guide|safe }}
            </div>
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endfor %}
</body>
</html>
"""
